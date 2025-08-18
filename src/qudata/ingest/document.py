"""
Document format extraction module for LLMDataForge.

This module provides extraction capabilities for DOCX files using python-docx,
with support for embedded objects and semantic structure preservation.
"""

import os
import re
from typing import Dict, List, Optional, Any, Tuple
from pathlib import Path

# Optional dependency for DOCX processing
try:
    import docx
    from docx.document import Document as DocxDocument
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from ..models import (
    BaseExtractor, ExtractedContent, FileMetadata, DocumentStructure,
    ProcessingError, ErrorSeverity, TableData, ImageData
)


class DocumentExtractor(BaseExtractor):
    """
    Extractor for DOCX files using python-docx.
    
    Handles text extraction, table detection, image metadata extraction,
    and structure analysis with robust error handling for corrupted files.
    """
    
    SUPPORTED_FORMATS = {'docx', 'doc'}
    
    def __init__(self, config: Dict[str, Any] = None):
        """
        Initialize the DocumentExtractor.
        
        Args:
            config: Configuration dictionary with options like:
                - max_file_size: Maximum file size to process (default: 100MB)
                - extract_tables: Whether to extract tables (default: True)
                - extract_images: Whether to extract image metadata (default: True)
                - preserve_formatting: Whether to preserve text formatting (default: False)
                - include_headers_footers: Whether to include headers/footers (default: False)
        """
        super().__init__(config)
        
        if not HAS_DOCX:
            raise ImportError(
                "python-docx is required for DOCX extraction. "
                "Install it with: pip install python-docx"
            )
        
        self.max_file_size = self.config.get('max_file_size', 100 * 1024 * 1024)  # 100MB
        self.extract_tables = self.config.get('extract_tables', True)
        self.extract_images = self.config.get('extract_images', True)
        self.preserve_formatting = self.config.get('preserve_formatting', False)
        self.include_headers_footers = self.config.get('include_headers_footers', False)
    
    def extract(self, file_path: str) -> ExtractedContent:
        """
        Extract content from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file to extract content from
            
        Returns:
            ExtractedContent object containing the extracted content and metadata
            
        Raises:
            ProcessingError: If extraction fails
        """
        try:
            # Validate file first
            self.validate_file(file_path)
            
            # Get basic file metadata
            file_metadata = self.get_metadata(file_path)
            
            # Check file size
            if file_metadata.size_bytes > self.max_file_size:
                raise self.create_processing_error(
                    stage="extraction",
                    error_type="FileTooLarge",
                    message=f"File size ({file_metadata.size_bytes} bytes) exceeds maximum ({self.max_file_size} bytes)",
                    severity=ErrorSeverity.HIGH
                )
            
            # Extract content from DOCX
            content, structure, tables, images = self._extract_docx_content(file_path)
            
            # Create extracted content object
            extracted = ExtractedContent(content, file_metadata)
            extracted.structure = structure
            extracted.tables = tables
            extracted.images = images
            
            return extracted
            
        except ProcessingError:
            raise
        except Exception as e:
            raise self.create_processing_error(
                stage="extraction",
                error_type="DOCXExtractionError",
                message=f"Failed to extract content from DOCX {file_path}: {str(e)}",
                severity=ErrorSeverity.HIGH,
                stack_trace=str(e)
            )
    
    def supports_format(self, file_type: str) -> bool:
        """
        Check if this extractor supports the given file type.
        
        Args:
            file_type: The file type to check
            
        Returns:
            True if the extractor supports this file type, False otherwise
        """
        return file_type.lower() in self.SUPPORTED_FORMATS
    
    def _extract_docx_content(self, file_path: str) -> Tuple[str, DocumentStructure, List[TableData], List[ImageData]]:
        """
        Extract content, structure, tables, and images from DOCX.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (content, structure, tables, images)
            
        Raises:
            ProcessingError: If DOCX cannot be processed
        """
        try:
            # Open the DOCX document
            doc = docx.Document(file_path)
            
            # Extract main document content
            content_parts = []
            tables = []
            images = []
            headings = []
            paragraphs = 0
            lists = 0
            
            # Process document elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    paragraph = DocxParagraph(element, doc)
                    para_text = paragraph.text.strip()
                    
                    if para_text:
                        # Check if this is a heading
                        is_heading = self._is_heading_paragraph(paragraph)
                        if is_heading:
                            headings.append(para_text)
                        
                        # Check if this is a list item
                        if self._is_list_paragraph(paragraph):
                            lists += 1
                        
                        # Add paragraph text
                        if self.preserve_formatting:
                            formatted_text = self._extract_formatted_text(paragraph)
                            content_parts.append(formatted_text)
                        else:
                            content_parts.append(para_text)
                        
                        # Only count as paragraph if it's not a heading
                        if not is_heading:
                            paragraphs += 1
                
                elif element.tag.endswith('tbl') and self.extract_tables:  # Table
                    table = DocxTable(element, doc)
                    table_data = self._extract_table_data(table, len(tables) + 1)
                    if table_data:
                        tables.append(table_data)
                        # Add table representation to content
                        content_parts.append(self._table_to_text(table_data))
            
            # Extract images if enabled
            if self.extract_images:
                images = self._extract_image_metadata(doc)
            
            # Include headers and footers if requested
            if self.include_headers_footers:
                header_footer_content = self._extract_headers_footers(doc)
                if header_footer_content:
                    content_parts.insert(0, header_footer_content)
            
            # Combine all content
            full_content = '\n\n'.join(content_parts) if content_parts else ""
            
            # Create document structure
            structure = DocumentStructure(
                headings=headings,
                paragraphs=paragraphs,
                tables=len(tables),
                images=len(images),
                code_blocks=0,  # DOCX doesn't typically have code blocks
                lists=lists
            )
            
            return full_content, structure, tables, images
            
        except FileNotFoundError:
            raise self.create_processing_error(
                stage="extraction",
                error_type="FileNotFound",
                message=f"DOCX file not found: {file_path}",
                severity=ErrorSeverity.HIGH
            )
        except PermissionError:
            raise self.create_processing_error(
                stage="extraction",
                error_type="PermissionError",
                message=f"Permission denied accessing DOCX file: {file_path}",
                severity=ErrorSeverity.HIGH
            )
        except Exception as e:
            # Handle various DOCX parsing errors
            error_message = str(e)
            if "corrupt" in error_message.lower() or "invalid" in error_message.lower():
                raise self.create_processing_error(
                    stage="extraction",
                    error_type="CorruptedDOCX",
                    message=f"DOCX file {file_path} appears to be corrupted: {str(e)}",
                    severity=ErrorSeverity.HIGH,
                    stack_trace=str(e)
                )
            elif "password" in error_message.lower() or "encrypted" in error_message.lower():
                raise self.create_processing_error(
                    stage="extraction",
                    error_type="EncryptedDOCX",
                    message=f"DOCX file {file_path} is password protected: {str(e)}",
                    severity=ErrorSeverity.HIGH,
                    stack_trace=str(e)
                )
            else:
                # Re-raise as generic extraction error
                raise
    
    def _is_heading_paragraph(self, paragraph: DocxParagraph) -> bool:
        """
        Determine if a paragraph is a heading.
        
        Args:
            paragraph: The paragraph to check
            
        Returns:
            True if paragraph is a heading
        """
        # Check if paragraph has a heading style
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name or 'title' in style_name:
                return True
        
        # Check formatting characteristics
        if paragraph.runs:
            first_run = paragraph.runs[0]
            # Check if text is bold and/or larger font
            if (hasattr(first_run, 'bold') and first_run.bold) or \
               (hasattr(first_run, 'font') and first_run.font.size and 
                first_run.font.size.pt > 12):
                # Additional check: heading-like text (short, no ending punctuation)
                text = paragraph.text.strip()
                if len(text) < 100 and not text.endswith(('.', '!', '?')):
                    return True
        
        return False
    
    def _is_list_paragraph(self, paragraph: DocxParagraph) -> bool:
        """
        Determine if a paragraph is part of a list.
        
        Args:
            paragraph: The paragraph to check
            
        Returns:
            True if paragraph is a list item
        """
        # Check if paragraph has list formatting
        if hasattr(paragraph, '_element') and paragraph._element is not None:
            # Check for numbering properties
            pPr = paragraph._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
            if pPr is not None:
                numPr = pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    return True
        
        # Fallback: check text patterns
        text = paragraph.text.strip()
        if text:
            # Check for bullet or number patterns
            if re.match(r'^\s*[-•*]\s+', text) or re.match(r'^\s*\d+[\.\)]\s+', text):
                return True
        
        return False
    
    def _extract_formatted_text(self, paragraph: DocxParagraph) -> str:
        """
        Extract text with basic formatting preservation.
        
        Args:
            paragraph: The paragraph to extract text from
            
        Returns:
            Text with basic formatting markers
        """
        formatted_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Add basic formatting markers
            if hasattr(run, 'bold') and run.bold:
                text = f"**{text}**"
            if hasattr(run, 'italic') and run.italic:
                text = f"*{text}*"
            if hasattr(run, 'underline') and run.underline:
                text = f"_{text}_"
            
            formatted_parts.append(text)
        
        return ''.join(formatted_parts)
    
    def _extract_table_data(self, table: DocxTable, table_num: int) -> Optional[TableData]:
        """
        Extract data from a DOCX table.
        
        Args:
            table: The DOCX table object
            table_num: Table number for reference
            
        Returns:
            TableData object or None if table is empty
        """
        try:
            rows = []
            headers = []
            
            # Extract table rows
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    # Extract text from all paragraphs in the cell
                    cell_text = '\n'.join(paragraph.text.strip() for paragraph in cell.paragraphs)
                    row_data.append(cell_text.strip())
                
                if row_idx == 0:
                    # First row as headers
                    headers = row_data
                else:
                    rows.append(row_data)
            
            # Only return table if it has meaningful content
            if headers and any(header.strip() for header in headers):
                return TableData(
                    headers=headers,
                    rows=rows,
                    caption=f"Table {table_num}"
                )
            
            return None
            
        except Exception as e:
            # Log table extraction errors but don't fail the entire extraction
            print(f"Warning: Error extracting table {table_num}: {str(e)}")
            return None
    
    def _table_to_text(self, table_data: TableData) -> str:
        """
        Convert table data to text representation.
        
        Args:
            table_data: The table data to convert
            
        Returns:
            Text representation of the table
        """
        lines = []
        
        if table_data.caption:
            lines.append(f"[{table_data.caption}]")
        
        # Add headers
        if table_data.headers:
            lines.append(" | ".join(table_data.headers))
            lines.append(" | ".join(["---"] * len(table_data.headers)))
        
        # Add rows
        for row in table_data.rows:
            lines.append(" | ".join(row))
        
        return "\n".join(lines)
    
    def _extract_image_metadata(self, doc: DocxDocument) -> List[ImageData]:
        """
        Extract image metadata from DOCX document.
        
        Args:
            doc: The DOCX document object
            
        Returns:
            List of ImageData objects
        """
        images = []
        
        try:
            # Get document relationships to find images
            if hasattr(doc, 'part') and hasattr(doc.part, 'rels'):
                image_count = 0
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        image_count += 1
                        # Create image metadata
                        image_data = ImageData(
                            path=f"embedded_image_{image_count}",
                            caption=f"Embedded image {image_count}",
                            alt_text=f"Image extracted from DOCX document"
                        )
                        images.append(image_data)
                        
        except Exception as e:
            # Log image extraction errors but don't fail the entire extraction
            print(f"Warning: Error extracting image metadata: {str(e)}")
        
        return images
    
    def _extract_headers_footers(self, doc: DocxDocument) -> str:
        """
        Extract headers and footers from DOCX document.
        
        Args:
            doc: The DOCX document object
            
        Returns:
            Combined headers and footers text
        """
        header_footer_parts = []
        
        try:
            # Extract headers
            for section in doc.sections:
                if section.header:
                    header_text = '\n'.join(paragraph.text.strip() 
                                          for paragraph in section.header.paragraphs 
                                          if paragraph.text.strip())
                    if header_text:
                        header_footer_parts.append(f"[Header: {header_text}]")
                
                if section.footer:
                    footer_text = '\n'.join(paragraph.text.strip() 
                                          for paragraph in section.footer.paragraphs 
                                          if paragraph.text.strip())
                    if footer_text:
                        header_footer_parts.append(f"[Footer: {footer_text}]")
                        
        except Exception as e:
            # Log header/footer extraction errors but don't fail
            print(f"Warning: Error extracting headers/footers: {str(e)}")
        
        return '\n'.join(header_footer_parts)