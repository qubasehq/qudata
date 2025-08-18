"""
Unit tests for DocumentExtractor (DOCX files).

Tests the extraction of content, tables, images, and structure from DOCX files
with various formatting and embedded objects.
"""

import os
import tempfile
import unittest
from unittest.mock import Mock, patch, MagicMock
from pathlib import Path

# Test if python-docx is available
try:
    import docx
    from docx.shared import Inches
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from src.qudata.ingest.document import DocumentExtractor
from src.qudata.models import ProcessingError, ErrorSeverity, DocumentStructure, TableData, ImageData


class TestDocumentExtractor(unittest.TestCase):
    """Test cases for DocumentExtractor."""
    
    def setUp(self):
        """Set up test fixtures."""
        self.config = {
            'max_file_size': 10 * 1024 * 1024,  # 10MB
            'extract_tables': True,
            'extract_images': True,
            'preserve_formatting': False,
            'include_headers_footers': False
        }
        
        if HAS_DOCX:
            self.extractor = DocumentExtractor(self.config)
        else:
            # Skip tests if python-docx is not available
            self.skipTest("python-docx not available")
    
    def test_supports_format(self):
        """Test file format support detection."""
        self.assertTrue(self.extractor.supports_format('docx'))
        self.assertTrue(self.extractor.supports_format('doc'))
        self.assertTrue(self.extractor.supports_format('DOCX'))
        self.assertFalse(self.extractor.supports_format('pdf'))
        self.assertFalse(self.extractor.supports_format('txt'))
    
    def test_initialization_without_docx(self):
        """Test initialization fails gracefully without python-docx."""
        with patch('src.qudata.ingest.document.HAS_DOCX', False):
            with self.assertRaises(ImportError) as context:
                DocumentExtractor()
            self.assertIn("python-docx is required", str(context.exception))
    
    def test_extract_simple_document(self):
        """Test extraction from a simple DOCX document."""
        # Create a simple test document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc = docx.Document()
            doc.add_heading('Test Document', 0)
            doc.add_paragraph('This is a test paragraph.')
            doc.add_paragraph('This is another paragraph with some content.')
            doc.save(tmp_file.name)
            
            try:
                # Extract content
                result = self.extractor.extract(tmp_file.name)
                
                # Verify basic properties
                self.assertIsNotNone(result)
                self.assertIsNotNone(result.content)
                self.assertIn('Test Document', result.content)
                self.assertIn('test paragraph', result.content)
                
                # Verify structure
                self.assertIsNotNone(result.structure)
                self.assertEqual(len(result.structure.headings), 1)
                self.assertEqual(result.structure.headings[0], 'Test Document')
                self.assertEqual(result.structure.paragraphs, 2)
                
                # Verify metadata
                self.assertEqual(result.metadata.file_type, 'docx')
                self.assertGreater(result.metadata.size_bytes, 0)
                
            finally:
                os.unlink(tmp_file.name)
    
    def test_extract_document_with_tables(self):
        """Test extraction from DOCX document with tables."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc = docx.Document()
            doc.add_heading('Document with Table', 0)
            doc.add_paragraph('This document contains a table.')
            
            # Add a table
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Name'
            header_cells[1].text = 'Age'
            header_cells[2].text = 'City'
            
            # Add data rows
            row1_cells = table.rows[1].cells
            row1_cells[0].text = 'John'
            row1_cells[1].text = '25'
            row1_cells[2].text = 'New York'
            
            row2_cells = table.rows[2].cells
            row2_cells[0].text = 'Jane'
            row2_cells[1].text = '30'
            row2_cells[2].text = 'London'
            
            doc.save(tmp_file.name)
            
            try:
                # Extract content
                result = self.extractor.extract(tmp_file.name)
                
                # Verify table extraction
                self.assertEqual(len(result.tables), 1)
                table_data = result.tables[0]
                
                self.assertEqual(table_data.headers, ['Name', 'Age', 'City'])
                self.assertEqual(len(table_data.rows), 2)
                self.assertEqual(table_data.rows[0], ['John', '25', 'New York'])
                self.assertEqual(table_data.rows[1], ['Jane', '30', 'London'])
                
                # Verify structure includes table
                self.assertEqual(result.structure.tables, 1)
                
                # Verify table is included in content
                self.assertIn('Name | Age | City', result.content)
                
            finally:
                os.unlink(tmp_file.name)
    
    def test_extract_document_with_formatting(self):
        """Test extraction from DOCX document with various formatting."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc = docx.Document()
            
            # Add various heading levels
            doc.add_heading('Main Title', 0)
            doc.add_heading('Section 1', 1)
            doc.add_heading('Subsection 1.1', 2)
            
            # Add paragraph with formatting
            p = doc.add_paragraph()
            run = p.add_run('This is bold text')
            run.bold = True
            p.add_run(' and this is ')
            run = p.add_run('italic text')
            run.italic = True
            p.add_run('.')
            
            # Add a list (simulated with bullet points)
            doc.add_paragraph('• First item', style='List Bullet')
            doc.add_paragraph('• Second item', style='List Bullet')
            doc.add_paragraph('• Third item', style='List Bullet')
            
            doc.save(tmp_file.name)
            
            try:
                # Test without formatting preservation
                result = self.extractor.extract(tmp_file.name)
                
                # Verify headings are detected
                self.assertEqual(len(result.structure.headings), 3)
                self.assertIn('Main Title', result.structure.headings)
                self.assertIn('Section 1', result.structure.headings)
                self.assertIn('Subsection 1.1', result.structure.headings)
                
                # Verify lists are detected
                self.assertGreaterEqual(result.structure.lists, 3)
                
                # Test with formatting preservation
                self.extractor.config['preserve_formatting'] = True
                result_formatted = self.extractor.extract(tmp_file.name)
                
                # Should contain formatting markers
                self.assertIn('**', result_formatted.content)  # Bold markers
                self.assertIn('*', result_formatted.content)   # Italic markers
                
            finally:
                os.unlink(tmp_file.name)
    
    def test_extract_document_with_headers_footers(self):
        """Test extraction with headers and footers."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc = docx.Document()
            
            # Add content
            doc.add_heading('Main Content', 0)
            doc.add_paragraph('This is the main document content.')
            
            # Add header
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = "Document Header"
            
            # Add footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Document Footer"
            
            doc.save(tmp_file.name)
            
            try:
                # Test without headers/footers
                result = self.extractor.extract(tmp_file.name)
                self.assertNotIn('Document Header', result.content)
                self.assertNotIn('Document Footer', result.content)
                
                # Test with headers/footers
                self.extractor.config['include_headers_footers'] = True
                result_with_hf = self.extractor.extract(tmp_file.name)
                self.assertIn('Header: Document Header', result_with_hf.content)
                self.assertIn('Footer: Document Footer', result_with_hf.content)
                
            finally:
                os.unlink(tmp_file.name)
    
    def test_file_size_limit(self):
        """Test file size limit enforcement."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc = docx.Document()
            doc.add_paragraph('Test content')
            doc.save(tmp_file.name)
            
            try:
                # Set very small file size limit
                small_extractor = DocumentExtractor({'max_file_size': 100})  # 100 bytes
                
                with self.assertRaises(ProcessingError) as context:
                    small_extractor.extract(tmp_file.name)
                
                error = context.exception
                self.assertEqual(error.error_type, 'FileTooLarge')
                self.assertEqual(error.severity, ErrorSeverity.HIGH)
                
            finally:
                os.unlink(tmp_file.name)
    
    def test_nonexistent_file(self):
        """Test handling of nonexistent files."""
        with self.assertRaises(ProcessingError) as context:
            self.extractor.extract('/nonexistent/file.docx')
        
        error = context.exception
        self.assertEqual(error.error_type, 'FileNotFound')
        self.assertEqual(error.severity, ErrorSeverity.HIGH)
    
    def test_corrupted_file(self):
        """Test handling of corrupted DOCX files."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            # Write invalid content
            tmp_file.write(b'This is not a valid DOCX file')
            tmp_file.flush()
            
            try:
                with self.assertRaises(ProcessingError) as context:
                    self.extractor.extract(tmp_file.name)
                
                error = context.exception
                self.assertIn('DOCX', error.error_type)
                self.assertEqual(error.severity, ErrorSeverity.HIGH)
                
            finally:
                os.unlink(tmp_file.name)
    
    def test_empty_document(self):
        """Test extraction from empty DOCX document."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc = docx.Document()
            doc.save(tmp_file.name)
            
            try:
                result = self.extractor.extract(tmp_file.name)
                
                # Should handle empty document gracefully
                self.assertIsNotNone(result)
                self.assertEqual(result.content.strip(), '')
                self.assertEqual(len(result.structure.headings), 0)
                self.assertEqual(result.structure.paragraphs, 0)
                self.assertEqual(len(result.tables), 0)
                
            finally:
                os.unlink(tmp_file.name)
    
    def test_complex_document_structure(self):
        """Test extraction from complex document with mixed content."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc = docx.Document()
            
            # Add title
            doc.add_heading('Complex Document', 0)
            
            # Add introduction
            doc.add_paragraph('This document contains various elements to test extraction.')
            
            # Add section with subsections
            doc.add_heading('Section 1: Tables', 1)
            doc.add_paragraph('This section contains a table.')
            
            # Add table
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = 'Header 1'
            table.rows[0].cells[1].text = 'Header 2'
            table.rows[1].cells[0].text = 'Data 1'
            table.rows[1].cells[1].text = 'Data 2'
            
            # Add another section
            doc.add_heading('Section 2: Lists', 1)
            doc.add_paragraph('This section contains lists.')
            doc.add_paragraph('• Item 1', style='List Bullet')
            doc.add_paragraph('• Item 2', style='List Bullet')
            
            # Add subsection
            doc.add_heading('Subsection 2.1', 2)
            doc.add_paragraph('This is a subsection with more content.')
            
            doc.save(tmp_file.name)
            
            try:
                result = self.extractor.extract(tmp_file.name)
                
                # Verify comprehensive structure detection
                self.assertEqual(len(result.structure.headings), 4)
                self.assertGreater(result.structure.paragraphs, 3)
                self.assertEqual(result.structure.tables, 1)
                self.assertGreaterEqual(result.structure.lists, 2)
                
                # Verify content includes all sections
                self.assertIn('Complex Document', result.content)
                self.assertIn('Section 1: Tables', result.content)
                self.assertIn('Section 2: Lists', result.content)
                self.assertIn('Subsection 2.1', result.content)
                
                # Verify table content
                self.assertEqual(len(result.tables), 1)
                self.assertEqual(result.tables[0].headers, ['Header 1', 'Header 2'])
                
            finally:
                os.unlink(tmp_file.name)
    
    @patch('src.qudata.ingest.document.docx.Document')
    def test_extraction_with_mock_error(self, mock_docx):
        """Test error handling during extraction."""
        # Mock docx.Document to raise an exception
        mock_docx.side_effect = Exception("Mock extraction error")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(b'dummy content')
            tmp_file.flush()
            
            try:
                with self.assertRaises(ProcessingError) as context:
                    self.extractor.extract(tmp_file.name)
                
                error = context.exception
                self.assertEqual(error.error_type, 'DOCXExtractionError')
                self.assertIn('Mock extraction error', error.message)
                
            finally:
                os.unlink(tmp_file.name)


if __name__ == '__main__':
    unittest.main()